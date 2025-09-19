import os
import numpy as np
import json
import subprocess
import sys
from mod.crear_latex import agregar_grafica_latex
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

size_titulo = 20
size_subtitulo = 14

def agregar_texto(doc, texto, tamano=12, negrita=False, alineacion=WD_PARAGRAPH_ALIGNMENT.LEFT):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(tamano)
    run.font.bold = negrita
    p.alignment = alineacion
    return p

def crear_archivo_word(tipo_grafica, eje_x, ejes_y, base_dir):
    carpeta = os.path.join(base_dir, "temporales")
    resultados_dir = os.path.join(base_dir, "resultados")
    datos_path = os.path.join(carpeta, "datos.dat")
    config_path = os.path.join(carpeta, "configuracion_grafica.json")
    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    datos = np.loadtxt(datos_path, delimiter="\t")
    idx_x = -1 if eje_x == "X0" else int(eje_x[1:]) - 1
    idx_ys = [int(y[1:]) - 1 for y in ejes_y]
    nombre_xs = config.get("nombre_x", [])

    doc = Document()
    agregar_texto(
        doc,
        "Tratamiento de Datos y Ajuste de Curvas",
        tamano=size_titulo,
        negrita=True,
        alineacion=WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    for i, idx_y in enumerate(idx_ys):
        nombre_y = nombre_xs[i] if i < len(nombre_xs) and nombre_xs[i] else f"Y{i+1}"
        if tipo_grafica.lower() == "lineal":
            agregar_texto(
                doc,
                f"Ajuste Lineal para {nombre_y}",
                tamano=size_subtitulo,
                negrita=True,
                alineacion=WD_PARAGRAPH_ALIGNMENT.LEFT
            )
        elif tipo_grafica.lower() == "exponencial":
            agregar_texto(
                doc,
                f"Ajuste Exponencial para {nombre_y}",
                tamano=size_subtitulo,
                negrita=True,
                alineacion=WD_PARAGRAPH_ALIGNMENT.LEFT
            )
        tabla = crear_tabla_word(doc, datos, idx_x, idx_y, tipo_grafica)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        agregar_ecuaciones_word(doc, datos, idx_x, idx_y, tipo_grafica)

    img_path = crear_grafica_latex_y_png(datos, idx_x, idx_ys, tipo_grafica, config, carpeta)
    if img_path and os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    os.makedirs(carpeta, exist_ok=True)
    ruta_docx = os.path.join(carpeta, "analisis.docx")
    doc.save(ruta_docx)
    os.makedirs(resultados_dir, exist_ok=True)
    docx_destino = os.path.join(resultados_dir, "analisis.docx")
    try:
        if os.path.exists(docx_destino):
            os.remove(docx_destino)
        os.replace(ruta_docx, docx_destino)
    except Exception as e:
        print("No se pudo mover el DOCX a la carpeta Resultados:", e)
        docx_destino = ruta_docx
    try:
        if sys.platform.startswith("win"):
            os.startfile(docx_destino)
        elif sys.platform.startswith("darwin"):
            subprocess.run(["open", docx_destino])
        else:
            subprocess.run(["xdg-open", docx_destino])
    except Exception:
        pass
    return docx_destino

def crear_tabla_word(doc, datos, idx_x, idx_y, tipo_grafica):
    if tipo_grafica.lower() == "exponencial":
        headers = [r"$x_i$", r"$y_i$", r"$\log x_i$", r"$\log y_i$", r"$\log x_i \log y_i$", r"$(\log x_i)^2$"]
        rows = []
        suma_x = suma_y = suma_logx = suma_logy = suma_logxlogy = suma_logx2 = 0
        x_vals = np.array([fila+1 if idx_x == -1 else datos[fila, idx_x] for fila in range(datos.shape[0])])
        y_vals = datos[:, idx_y]
        for x, y in zip(x_vals, y_vals):
            logx = np.log10(x)
            logy = np.log10(y)
            logxlogy = logx * logy
            logx2 = logx ** 2
            suma_x += x
            suma_y += y
            suma_logx += logx
            suma_logy += logy
            suma_logxlogy += logxlogy
            suma_logx2 += logx2
            rows.append([f"{x:.4f}", f"{y:.4f}", f"{logx:.10f}", f"{logy:.10f}", f"{logxlogy:.10f}", f"{logx2:.10f}"])
        sum_labels = [r"$\sum x_i$", r"$\sum y_i$", r"$\sum \log x_i$", r"$\sum \log y_i$", r"$\sum \log x_i \log y_i$", r"$\sum (\log x_i)^2$"]
        sum_row = [f"{suma_x:.4f}", f"{suma_y:.4f}", f"{suma_logx:.10f}", f"{suma_logy:.10f}", f"{suma_logxlogy:.10f}", f"{suma_logx2:.10f}"]
    else:
        headers = [r"$x_i$", r"$y_i$", r"$x_i y_i$", r"$x_i^2$"]
        rows = []
        suma_x = suma_y = suma_xy = suma_x2 = 0
        x_vals = np.array([fila+1 if idx_x == -1 else datos[fila, idx_x] for fila in range(datos.shape[0])])
        y_vals = datos[:, idx_y]
        for x, y in zip(x_vals, y_vals):
            xy = x * y
            x2 = x ** 2
            suma_x += x
            suma_y += y
            suma_xy += xy
            suma_x2 += x2
            rows.append([f"{x:.4f}", f"{y:.4f}", f"{xy:.4f}", f"{x2:.4f}"])
        sum_labels = [r"$\sum x_i$", r"$\sum y_i$", r"$\sum x_i y_i$", r"$\sum x_i^2$"]
        sum_row = [f"{suma_x:.4f}", f"{suma_y:.4f}", f"{suma_xy:.4f}", f"{suma_x2:.4f}"]

    table = doc.add_table(rows=1 + len(rows) + 2, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = h
        for paragraph in hdr_cells[j].paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i+1, j)
            cell.text = val
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sum_label_idx = len(rows) + 1
    for j, label in enumerate(sum_labels):
        cell = table.cell(sum_label_idx, j)
        cell.text = label
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sum_row_idx = len(rows) + 2
    for j, val in enumerate(sum_row):
        cell = table.cell(sum_row_idx, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return table

def agregar_ecuaciones_word(doc, datos, idx_x, idx_y, tipo_grafica):
    agregar_texto(doc, "Ecuaciones", tamano=size_subtitulo, negrita=True, alineacion=WD_PARAGRAPH_ALIGNMENT.LEFT)
    
    if idx_x == -1:
        x = np.arange(1, datos.shape[0]+1)
    else:
        x = datos[:, idx_x]
    y = datos[:, idx_y]
    if tipo_grafica.lower() == "lineal":
        p = len(x)
        suma_x = np.sum(x)
        suma_y = np.sum(y)
        suma_xy = np.sum(x * y)
        suma_x2 = np.sum(x ** 2)
        denominador = (p * suma_x2 - suma_x ** 2.0)
        if denominador != 0:
            m = (p * suma_xy - suma_x * suma_y) / denominador
            b = (suma_x2 * suma_y - suma_x * suma_xy) / denominador
            ecuaciones = [
                r"Ecuación general:",
                r"$y = m x + b$",
                r"Desarrollo de $m$:",
                r"$m = \frac{p \sum x_i y_i - \sum x_i \sum y_i}{p \sum x_i^2 - (\sum x_i)^2}$",
                r"$m = \frac{%d \cdot %.4f - %.4f \cdot %.4f}{%d \cdot %.4f - (%.4f)^2}$" % (
                    p, suma_xy, suma_x, suma_y, p, suma_x2, suma_x
                ),
                r"$m = \frac{%.4f}{%.4f} = %.4f$" % (
                    (p * suma_xy - suma_x * suma_y),
                    denominador,
                    m
                ),
                r"Desarrollo de $b$:",
                r"$b = \frac{\sum x_i^2 \sum y_i - \sum x_i \sum x_i y_i}{p \sum x_i^2 - (\sum x_i)^2}$",
                r"$b = \frac{%.4f \cdot %.4f - %.4f \cdot %.4f}{%d \cdot %.4f - (%.4f)^2}$" % (
                    suma_x2, suma_y, suma_x, suma_xy, p, suma_x2, suma_x
                ),
                r"$b = \frac{%.4f}{%.4f} = %.4f$" % (
                    (suma_x2 * suma_y - suma_x * suma_xy),
                    denominador,
                    b
                ),
                r"Ecuación resultante:",
                r"$y = %.4fx %s%.4f$" % (m, "+" if b >= 0 else "", b)
            ]
            for eq in ecuaciones:
                p = doc.add_paragraph()
                p.add_run(eq)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif tipo_grafica.lower() == "exponencial":
        p = len(x)
        log_x = np.log10(x)
        log_y = np.log10(y)
        suma_log_x = np.sum(log_x)
        suma_log_y = np.sum(log_y)
        suma_log_x_log_y = np.sum(log_x * log_y)
        suma_log_x2 = np.sum(log_x ** 2)
        denominador = (p * suma_log_x2 - suma_log_x ** 2.0)
        if denominador != 0:
            m = (p * suma_log_x_log_y - suma_log_x * suma_log_y) / denominador
            b = (suma_log_x2 * suma_log_y - suma_log_x * suma_log_x_log_y) / denominador
            ecuaciones = [
                r"Ecuación general:",
                r"$y = 10^{b} x^{m}$",
                "",
                r"Desarrollo de $m$:",
                r"$m = \frac{p \sum \log{x_i} \log{y_i} - \sum \log{x_i} \sum \log{y_i}}{p \sum \log{x_i}^2 - (\sum \log{x_i})^2}$",
                r"$m = \frac{%d \cdot %.4f - %.4f \cdot %.4f}{%d \cdot %.4f - (%.4f)^2}$" % (
                    p, suma_log_x_log_y, suma_log_x, suma_log_y, p, suma_log_x2, suma_log_x
                ),
                r"$m = \frac{%.4f}{%.4f} = %.4f$" % (
                    (p * suma_log_x_log_y - suma_log_x * suma_log_y),
                    denominador,
                    m
                ),
                "",
                r"Desarrollo de $b$:",
                r"$b = \frac{\sum \log{x_i}^2 \sum \log{y_i} - \sum \log{x_i} \sum \log{x_i} \log{y_i}}{p \sum \log{x_i}^2 - (\sum \log{x_i})^2}$",
                r"$b = \frac{%.4f \cdot %.4f - %.4f \cdot %.4f}{%d \cdot %.4f - (%.4f)^2}$" % (
                    suma_log_x2, suma_log_y, suma_log_x, suma_log_x_log_y, p, suma_log_x2, suma_log_x
                ),
                r"$b = \frac{%.4f}{%.4f} = %.4f$" % (
                    (suma_log_x2 * suma_log_y - suma_log_x * suma_log_x_log_y),
                    denominador,
                    b
                ),
                "",
                r"Ecuación resultante:",
                r"$y = 10^{%.4f} x^{%.4f}$" % (b, m)
            ]
            for eq in ecuaciones:
                p = doc.add_paragraph()
                p.add_run(eq)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def crear_grafica_latex_y_png(datos, idx_x, idx_ys, tipo_grafica, config, carpeta):
    contenido = [
        r"\documentclass{standalone}",
        r"\usepackage{pgfplots}",
        r"\pgfplotsset{compat=1.18}",
        r"\begin{document}",
    ]
    contenido += agregar_grafica_latex(datos, idx_x, idx_ys, tipo_grafica, config,True)
    contenido.append(r"\end{document}")

    tex_path = os.path.abspath(os.path.join(carpeta, "grafica_temp.tex"))
    pdf_path = os.path.abspath(os.path.join(carpeta, "grafica_temp.pdf"))
    png_path = os.path.abspath(os.path.join(carpeta, "grafica_temp.png"))

    with open(tex_path, "w", encoding="utf-8") as f:
        f.write("\n".join(contenido))

    try:
        import shutil
        if shutil.which("pdflatex") is None:
            return None

        # Compilar el archivo tex a pdf
        subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", tex_path],
            cwd=carpeta,
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )

        # Mejorar calidad: usar 300 DPI
        if shutil.which("pdftoppm"):
            subprocess.run(
                ["pdftoppm", "-png", "-singlefile", "-rx", "300", "-ry", "300", pdf_path, os.path.splitext(png_path)[0]],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
        elif shutil.which("convert"):
            subprocess.run(
                ["convert", "-density", "300", pdf_path, png_path],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
        else:
            return None

        # Limpiar archivos temporales
        for ext in [".aux", ".log", ".pdf", ".tex"]:
            temp_file = os.path.join(carpeta, f"grafica_temp{ext}")
            if os.path.exists(temp_file) and ext != ".png":
                try:
                    os.remove(temp_file)
                except Exception:
                    pass

        return png_path if os.path.exists(png_path) else None

    except Exception:
        return None

